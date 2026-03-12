#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word转换器

将Markdown文档转换为编排良好的Word格式
支持中文排版、表格、列表、代码块等
"""

import os
import sys
import re
import argparse
from pathlib import Path
from typing import Optional, List, Dict, Any

# 尝试导入必需的库
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    print("⚠️  警告: python-docx 未安装")
    print("请运行: pip install python-docx")


class MarkdownToWord:
    """Markdown转Word转换器"""
    
    def __init__(
        self,
        title: Optional[str] = None,
        author: Optional[str] = None,
        font_name: str = "Microsoft YaHei",
        font_size: int = 11
    ):
        """
        初始化转换器
        
        Args:
            title: 文档标题
            author: 作者信息
            font_name: 字体名称
            font_size: 正文字号
        """
        if not HAS_PYTHON_DOCX:
            raise RuntimeError("请先安装依赖库: pip install python-docx")
        
        self.title = title
        self.author = author
        self.font_name = font_name
        self.font_size = font_size
        
        # 创建文档
        self.doc = Document()
        
        # 设置文档默认字体
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        
        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        
        # 设置段落格式 - 更紧凑
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # 单倍行距
        paragraph_format.space_after = Pt(3)  # 段后间距减小
        paragraph_format.space_before = Pt(0)  # 段前间距
        
        # 设置标题样式
        for i in range(1, 5):
            heading_style = self.doc.styles[f'Heading {i}']
            font = heading_style.font
            font.name = self.font_name
            font.bold = True
            font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色，更正式
            
            # 设置中文字体
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            
            # 设置字号 - 更紧凑
            sizes = {1: 18, 2: 14, 3: 12, 4: 11}
            font.size = Pt(sizes[i])
            
            # 设置段落间距
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    def _add_heading(self, text: str, level: int):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        return heading
    
    def _add_paragraph(self, text: str, bold: bool = False, indent: bool = False):
        """添加段落"""
        para = self.doc.add_paragraph()
        
        # 处理文本中的加粗和斜体
        parts = []
        current = ""
        i = 0
        while i < len(text):
            # 处理加粗 **text**
            if text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                j = i + 2
                while j < len(text) and text[j:j+2] != '**':
                    j += 1
                parts.append(('bold', text[i+2:j]))
                i = j + 2
                continue
            # 处理斜体 *text*
            elif text[i] == '*' and (i == 0 or text[i-1] != '*'):
                if current:
                    parts.append(('normal', current))
                    current = ""
                j = i + 1
                while j < len(text) and text[j] != '*':
                    j += 1
                parts.append(('italic', text[i+1:j]))
                i = j + 1
                continue
            else:
                current += text[i]
                i += 1
        
        if current:
            parts.append(('normal', current))
        
        # 添加格式化文本
        for style_type, part_text in parts:
            run = para.add_run(part_text)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            
            if style_type == 'bold' or bold:
                run.bold = True
            elif style_type == 'italic':
                run.italic = True
        
        # 首行缩进
        if indent:
            para.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符
        
        return para
    
    def _add_bullet_list(self, items: List[str]):
        """添加无序列表"""
        for item in items:
            para = self.doc.add_paragraph(style='List Bullet')
            run = para.add_run(item)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
    
    def _add_numbered_list(self, items: List[str]):
        """添加有序列表"""
        for item in items:
            para = self.doc.add_paragraph(style='List Number')
            run = para.add_run(item)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
    
    def _add_table(self, rows: List[List[str]], header: bool = True):
        """添加表格"""
        if not rows:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格宽度
        for col in table.columns:
            col.width = Inches(1.5)
        
        # 填充表格
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text.strip()
                
                # 设置单元格字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.font_name
                        run.font.size = Pt(10)  # 表格字号稍小
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                
                # 表头加粗和背景色
                if header and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    # 设置背景色
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '003366')  # 深蓝色背景
                    cell._tc.get_or_add_tcPr().append(shading)
                elif i % 2 == 1:
                    # 隔行变色
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F5F5F5')
                    cell._tc.get_or_add_tcPr().append(shading)
        
        # 表格后添加空行
        self.doc.add_paragraph()
        
        return table
    
    def _add_code_block(self, code: str, language: str = ""):
        """添加代码块"""
        para = self.doc.add_paragraph()
        
        # 设置背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F4F4F4')
        para._p.get_or_add_pPr().append(shading)
        
        # 添加代码
        run = para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        
        return para
    
    def _add_blockquote(self, text: str):
        """添加引用块"""
        para = self.doc.add_paragraph()
        
        # 设置左侧边框
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:color'), '1A5490')
        pBdr.append(left)
        pPr.append(pBdr)
        
        # 设置背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F7FF')
        pPr.append(shading)
        
        # 添加文本
        run = para.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.italic = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        
        # 设置缩进
        para.paragraph_format.left_indent = Cm(1)
        
        return para
    
    def _parse_table(self, lines: List[str], start_idx: int) -> tuple:
        """解析Markdown表格"""
        rows = []
        i = start_idx
        
        while i < len(lines) and '|' in lines[i]:
            # 跳过分隔行
            if re.match(r'^\|[\s\-:]+\|$', lines[i].strip()):
                i += 1
                continue
            
            # 解析行
            cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            if cells:
                rows.append(cells)
            i += 1
        
        return rows, i
    
    def convert(self, md_file: str, output_file: str) -> str:
        """
        转换Markdown文件为Word
        
        Args:
            md_file: Markdown文件路径
            output_file: 输出Word文件路径
        
        Returns:
            输出文件路径
        """
        # 读取Markdown文件
        with open(md_file, "r", encoding="utf-8") as f:
            md_content = f.read()
        
        # 分割为行
        lines = md_content.split('\n')
        
        i = 0
        in_code_block = False
        code_lines = []
        code_language = ""
        
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.strip().startswith('```'):
                if in_code_block:
                    # 结束代码块
                    self._add_code_block('\n'.join(code_lines), code_language)
                    code_lines = []
                    in_code_block = False
                else:
                    # 开始代码块
                    in_code_block = True
                    code_language = line.strip()[3:].strip()
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            # 处理标题
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self._add_heading(text, level)
                i += 1
                continue
            
            # 处理表格
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                rows, i = self._parse_table(lines, i)
                if rows:
                    self._add_table(rows)
                continue
            
            # 处理无序列表
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                items = []
                while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                    items.append(lines[i].strip()[2:])
                    i += 1
                self._add_bullet_list(items)
                continue
            
            # 处理有序列表
            if re.match(r'^\d+\.\s+', line.strip()):
                items = []
                while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                    items.append(re.sub(r'^\d+\.\s+', '', lines[i].strip()))
                    i += 1
                self._add_numbered_list(items)
                continue
            
            # 处理引用块
            if line.strip().startswith('>'):
                text = line.strip()[1:].strip()
                self._add_blockquote(text)
                i += 1
                continue
            
            # 处理水平线
            if line.strip() in ['---', '***', '___']:
                para = self.doc.add_paragraph()
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                i += 1
                continue
            
            # 处理普通段落
            if line.strip():
                # 处理加粗、斜体等
                text = line.strip()
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 暂时简化处理
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                
                self._add_paragraph(text, indent=True)
            
            i += 1
        
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_file) if os.path.dirname(output_file) else '.', exist_ok=True)
        
        # 保存文档
        self.doc.save(output_file)
        
        return output_file


def generate_word(
    md_file: str,
    output_file: str,
    title: Optional[str] = None,
    author: Optional[str] = None,
    font_name: str = "Microsoft YaHei",
    font_size: int = 11
) -> str:
    """
    快捷函数：生成Word文档
    
    Args:
        md_file: Markdown文件路径
        output_file: 输出Word文件路径
        title: 文档标题（可选）
        author: 作者信息（可选）
        font_name: 字体名称
        font_size: 正文字号
    
    Returns:
        输出文件路径
    """
    converter = MarkdownToWord(
        title=title,
        author=author,
        font_name=font_name,
        font_size=font_size
    )
    return converter.convert(md_file, output_file)


def batch_convert(
    input_dir: str,
    output_dir: str,
    pattern: str = "*.md"
) -> List[str]:
    """
    批量转换目录下的Markdown文件
    
    Args:
        input_dir: 输入目录
        output_dir: 输出目录
        pattern: 文件匹配模式
    
    Returns:
        生成的Word文件列表
    """
    import glob
    
    md_files = glob.glob(os.path.join(input_dir, pattern))
    docx_files = []
    
    converter = MarkdownToWord()
    
    for md_file in md_files:
        base_name = os.path.splitext(os.path.basename(md_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}.docx")
        
        try:
            converter.convert(md_file, output_file)
            docx_files.append(output_file)
            print(f"[OK] 已转换: {md_file} -> {output_file}")
        except Exception as e:
            print(f"[ERROR] 转换失败 {md_file}: {str(e)}")
    
    return docx_files


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description="Markdown转Word转换器")
    parser.add_argument("--input", "-i", required=True, help="输入Markdown文件或目录")
    parser.add_argument("--output", "-o", required=True, help="输出Word文件或目录")
    parser.add_argument("--title", "-t", help="文档标题")
    parser.add_argument("--author", "-a", help="作者信息")
    parser.add_argument("--font", "-f", default="Microsoft YaHei", help="字体名称")
    parser.add_argument("--size", "-s", type=int, default=11, help="正文字号")
    parser.add_argument("--batch", "-b", action="store_true", help="批量转换模式")
    
    args = parser.parse_args()
    
    try:
        if args.batch:
            # 批量转换
            docx_files = batch_convert(args.input, args.output)
            print(f"\n[OK] 批量转换完成，共生成 {len(docx_files)} 个Word文件")
        else:
            # 单文件转换
            output = generate_word(
                md_file=args.input,
                output_file=args.output,
                title=args.title,
                author=args.author,
                font_name=args.font,
                font_size=args.size
            )
            print(f"[OK] Word生成成功: {output}")
            print(f"[INFO] 文件大小: {os.path.getsize(output) / 1024:.2f} KB")
    except Exception as e:
        print(f"[ERROR] 生成失败: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
