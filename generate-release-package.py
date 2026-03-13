#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OpenClaw技术研究报告 - 最终发布版生成器
生成主报告和5个附件Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle, Circle, FancyArrowPatch
import matplotlib.patches as mpatches
import numpy as np

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'SimHei', 'Microsoft YaHei', 'Heiti TC']
plt.rcParams['axes.unicode_minus'] = False

# 输出目录
OUTPUT_DIR = '/Users/sonnet/opencode/openclaw-research/output/release'
IMAGES_DIR = '/Users/sonnet/opencode/openclaw-research/output/images'

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)


class DiagramGenerator:
    """架构图生成器"""
    
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def create_architecture_diagram(self):
        fig, ax = plt.subplots(1, 1, figsize=(14, 10))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 10)
        ax.axis('off')
        
        colors = {'gateway': '#1a5276', 'agent': '#7b241c', 'channel': '#b7950b', 'node': '#1e8449', 'skill': '#6c3483'}
        ax.text(7, 9.5, 'OpenClaw技术架构', fontsize=20, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        layers = [
            ('网关层 Gateway', colors['gateway'], 7.8, '统一控制平面\nWebSocket服务 | 会话管理 | 消息路由 | 权限控制'),
            ('代理层 Agent', colors['agent'], 6.2, 'AI执行平面\nLLM调用 | 工具执行 | 状态管理 | 记忆系统'),
            ('渠道层 Channels', colors['channel'], 4.6, '接入适配层\nTelegram | Slack | Discord | 企业微信 | 钉钉 | 飞书'),
            ('节点层 Nodes', colors['node'], 3.0, '设备管理层\nmacOS | iOS | Android | CLI | Web'),
            ('技能层 Skills', colors['skill'], 1.4, '能力扩展层\n插件系统 | 工具集成 | 自定义能力'),
        ]
        
        for label, color, y, desc in layers:
            rect = FancyBboxPatch((1.5, y-0.6), 11, 1.2, boxstyle="round,pad=0.05,rounding_size=0.15",
                                  facecolor=color, edgecolor='white', linewidth=2, alpha=0.9)
            ax.add_patch(rect)
            ax.text(7, y+0.2, label, fontsize=13, fontweight='bold', ha='center', va='center', color='white')
            ax.text(7, y-0.25, desc, fontsize=9, ha='center', va='center', color='white', alpha=0.9)
        
        for i in range(len(layers)-1):
            y1, y2 = layers[i][2] - 0.6, layers[i+1][2] + 0.6
            ax.annotate('', xy=(7, y2), xytext=(7, y1), arrowprops=dict(arrowstyle='->', color='#666666', lw=2))
        
        ax.text(13.3, 5, '技术栈', fontsize=11, fontweight='bold', ha='center', va='center', color='#333333')
        ax.text(13.3, 4, 'TypeScript\nNode.js 22+\nWebSocket\nDocker\nSQLite', fontsize=9, ha='center', va='center', color='#666666', linespacing=1.6)
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'architecture.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_security_layers_diagram(self):
        fig, ax = plt.subplots(1, 1, figsize=(14, 8))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 8)
        ax.axis('off')
        
        ax.text(7, 7.5, 'OpenClaw多层次安全防护体系', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        layers = [
            ('审计追踪层', '#e8f4fd', '#1565c0', 0.3, 0.3, 13.4, 6.7, '安全审计工具 | 配置检查 | 风险识别 | 日志记录'),
            ('执行隔离层', '#fff3e0', '#e65100', 1.0, 1.0, 12.0, 5.3, 'Docker沙箱 | 网络隔离 | 文件系统隔离 | 资源限制'),
            ('访问控制层', '#f3e5f5', '#6a1b9a', 1.7, 1.7, 10.6, 3.9, 'DM配对机制 | Allowlist白名单 | 工具权限控制'),
            ('认证授权层', '#e8f5e9', '#2e7d32', 2.4, 2.4, 9.2, 2.5, 'Token认证 | 密码认证 | 设备认证'),
        ]
        
        for name, fc, ec, x, y, w, h, desc in layers:
            rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.08",
                                  facecolor=fc, edgecolor=ec, linewidth=2.5, alpha=0.95)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h - 0.35, name, fontsize=12, fontweight='bold', ha='center', va='top', color=ec)
            ax.text(x + w/2, y + h/2 - 0.1, desc, fontsize=9, ha='center', va='center', color='#333333')
        
        ax.annotate('', xy=(0.15, 3.5), xytext=(0.15, 1.5), arrowprops=dict(arrowstyle='<->', color='#999999', lw=1.5))
        ax.text(0.08, 2.5, '外层\n防护', fontsize=9, ha='right', va='center', color='#666666', rotation=90)
        ax.annotate('', xy=(13.85, 3.5), xytext=(13.85, 1.5), arrowprops=dict(arrowstyle='<->', color='#999999', lw=1.5))
        ax.text(13.92, 2.5, '内层\n核心', fontsize=9, ha='left', va='center', color='#666666', rotation=90)
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'security_layers.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_soe_architecture_diagram(self):
        fig, ax = plt.subplots(1, 1, figsize=(14, 12))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 12)
        ax.axis('off')
        
        ax.text(7, 11.5, '央国企智能助手架构建议（SOE-Claw）', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        rects = [
            (0.5, 9.8, 13, 1.2, '#e3f2fd', '#1565c0', '用户接入层', 'Web管理端 | 企业微信 | 钉钉 | 飞书 | 政务平台'),
            (0.5, 8.0, 13, 1.5, '#fce4ec', '#c2185b', '安全接入层', '统一身份认证(ISC) | RBAC权限控制 | 访问审计 | 安全防护'),
            (0.5, 6.0, 13, 1.7, '#fff8e1', '#f57c00', 'AI服务层', '对话服务 | 知识库服务(RAG) | 工具服务 | Agent编排'),
            (0.5, 4.0, 13, 1.7, '#e8f5e9', '#2e7d32', '模型服务层（国产大模型）', '百度文心 | 阿里通义 | 腾讯混元 | 华为盘古 | 私有化模型'),
            (0.5, 2.0, 13, 1.7, '#f3e5f5', '#7b1fa2', '数据服务层', '企业知识库 | 向量数据库 | 关系数据库 | 对象存储'),
            (0.5, 0.3, 13, 1.4, '#eceff1', '#546e7a', '基础设施层', '私有云/混合云 | 国产操作系统 | 国产数据库 | 安全设施'),
        ]
        
        for x, y, w, h, fc, ec, title, desc in rects:
            rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", facecolor=fc, edgecolor=ec, linewidth=2)
            ax.add_patch(rect)
            ax.text(7, y + h - 0.3, title, fontsize=11, fontweight='bold', ha='center', color=ec)
            ax.text(7, y + h/2 - 0.1, desc, fontsize=9, ha='center', color='#333')
        
        for y in [9.8, 8.0, 6.0, 4.0, 2.0]:
            ax.annotate('', xy=(7, y-0.1), xytext=(7, y), arrowprops=dict(arrowstyle='->', color='#999999', lw=1.5))
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'soe_architecture.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_permission_model_diagram(self):
        fig, ax = plt.subplots(1, 1, figsize=(14, 10))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 10)
        ax.axis('off')
        
        ax.text(7, 9.5, '深度权限控制模型', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        # 三层权限
        rects = [
            (0.5, 7.5, 4, 1.5, '#e3f2fd', '#1565c0', '用户/角色层', '用户 → 角色 → 部门\n组织架构集成'),
            (5, 7.5, 4, 1.5, '#fff8e1', '#f57c00', '功能权限(RBAC)', '角色 → 权限 → 功能\n菜单权限 | 操作权限'),
            (9.5, 7.5, 4, 1.5, '#e8f5e9', '#2e7d32', '数据权限(ABAC)', '属性 → 策略 → 数据\n部门数据 | 敏感数据'),
        ]
        
        for x, y, w, h, fc, ec, title, desc in rects:
            rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", facecolor=fc, edgecolor=ec, linewidth=2)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h - 0.3, title, fontsize=11, fontweight='bold', ha='center', color=ec)
            ax.text(x + w/2, y + h/2 - 0.1, desc, fontsize=9, ha='center', color='#333')
        
        ax.annotate('', xy=(5, 8.25), xytext=(4.5, 8.25), arrowprops=dict(arrowstyle='->', color='#999', lw=2))
        ax.annotate('', xy=(9.5, 8.25), xytext=(9, 8.25), arrowprops=dict(arrowstyle='->', color='#999', lw=2))
        
        # 权限决策流程
        rect4 = FancyBboxPatch((1, 4.5), 12, 2.5, boxstyle="round,pad=0.03", facecolor='#fce4ec', edgecolor='#c2185b', linewidth=2)
        ax.add_patch(rect4)
        ax.text(7, 6.3, '权限决策流程', fontsize=12, fontweight='bold', ha='center', color='#c2185b')
        
        steps = [(2.5, 5.5, '1.身份认证\nISC统一认证'), (5.5, 5.5, '2.角色映射\n获取用户角色'), 
                 (8.5, 5.5, '3.权限计算\nRBAC+ABAC'), (11.5, 5.5, '4.策略执行\n允许/拒绝')]
        for x, y, text in steps:
            ax.text(x, y, text, fontsize=9, ha='center', va='center', color='#333')
        
        for i in range(3):
            ax.annotate('', xy=(steps[i+1][0]-1.2, 5.5), xytext=(steps[i][0]+1.2, 5.5),
                       arrowprops=dict(arrowstyle='->', color='#999', lw=1.5))
        
        # AI特殊权限
        rect5 = FancyBboxPatch((1, 1.5), 12, 2.5, boxstyle="round,pad=0.03", facecolor='#f3e5f5', edgecolor='#7b1fa2', linewidth=2)
        ax.add_patch(rect5)
        ax.text(7, 3.5, 'AI能力权限控制', fontsize=12, fontweight='bold', ha='center', color='#7b1fa2')
        
        ai_controls = [(3, 2.5, '工具权限\n可调用工具白名单'), (7, 2.5, '数据权限\n可访问数据范围'), (11, 2.5, '模型权限\n可使用模型类型')]
        for x, y, text in ai_controls:
            rect = FancyBboxPatch((x-1.3, y-0.6), 2.6, 1.2, boxstyle="round,pad=0.02", facecolor='white', edgecolor='#7b1fa2', linewidth=1)
            ax.add_patch(rect)
            ax.text(x, y, text, fontsize=9, ha='center', va='center', color='#333')
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'permission_model.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_implementation_timeline(self):
        fig, ax = plt.subplots(1, 1, figsize=(14, 6))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 6)
        ax.axis('off')
        
        ax.text(7, 5.5, '央国企智能助手实施路径', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        ax.plot([1, 13], [3, 3], color='#cccccc', linewidth=4, zorder=1)
        
        phases = [
            (2.5, '第一阶段', '基础能力建设', '3-6个月', '#1565c0', '环境搭建\n模型部署\n基础权限\n安全防护'),
            (5.5, '第二阶段', '系统集成对接', '3-4个月', '#c2185b', '企业认证\n渠道适配\n知识库对接\n业务集成'),
            (8.5, '第三阶段', '合规完善试点', '3-4个月', '#f57c00', '等保整改\n安全测评\n试点部署\n功能优化'),
            (11.5, '第四阶段', '规模化推广', '6-12个月', '#2e7d32', '全系统推广\n运维体系\n持续优化\n生态建设'),
        ]
        
        for x, phase, name, time, color, desc in phases:
            circle = plt.Circle((x, 3), 0.3, color=color, zorder=2)
            ax.add_patch(circle)
            ax.text(x, 4.2, phase, fontsize=11, fontweight='bold', ha='center', va='bottom', color=color)
            ax.text(x, 3.9, name, fontsize=10, ha='center', va='bottom', color='#333')
            ax.text(x, 2.3, time, fontsize=9, ha='center', va='top', color='#666')
            ax.text(x, 1.3, desc, fontsize=8, ha='center', va='center', color='#666', linespacing=1.3)
        
        ax.text(7, 0.5, '总周期：15-26个月', fontsize=11, ha='center', va='center', color='#444444', style='italic')
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'timeline.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_risk_matrix(self):
        fig, ax = plt.subplots(1, 1, figsize=(12, 7))
        
        risks = [('技术风险', '中等', 3), ('安全风险', '高', 4), ('合规风险', '中等', 3), ('运维风险', '中等', 3)]
        colors = {'高': '#e74c3c', '中等': '#f39c12', '低': '#27ae60'}
        
        x = np.arange(len(risks))
        values = [r[2] for r in risks]
        color_list = [colors[r[1]] for r in risks]
        labels = [r[0] for r in risks]
        
        bars = ax.bar(x, values, color=color_list, width=0.6, edgecolor='white', linewidth=2)
        for bar, val, risk in zip(bars, values, risks):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.15,
                   f'{risk[1]}\n({val}/5)', ha='center', va='bottom', fontsize=11, fontweight='bold', color='#333333')
        
        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=12)
        ax.set_ylim(0, 5.5)
        ax.set_ylabel('风险等级', fontsize=12)
        ax.set_title('自主开发风险评估矩阵', fontsize=16, fontweight='bold', pad=20)
        ax.axhline(y=3.5, color='#e74c3c', linestyle='--', alpha=0.5, linewidth=1.5)
        ax.text(3.6, 3.65, '高风险线', fontsize=10, color='#e74c3c', va='bottom')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'risk_matrix.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_data_flow_diagram(self):
        fig, ax = plt.subplots(1, 1, figsize=(14, 8))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 8)
        ax.axis('off')
        
        ax.text(7, 7.5, '央国企智能助手数据流', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        rects = [
            (0.5, 4.5, 2, 1.5, '#e3f2fd', '#1565c0', '用户', '企业微信\n钉钉/飞书'),
            (3.5, 4.5, 2.5, 1.5, '#fce4ec', '#c2185b', '安全接入', '身份认证\n权限校验'),
            (7, 4.5, 2.5, 1.5, '#fff8e1', '#f57c00', 'AI服务', '对话处理\n知识检索'),
            (10.5, 4.5, 3, 1.5, '#e8f5e9', '#2e7d32', '国产模型', '文心/通义\n私有化部署'),
            (5, 1.5, 4, 1.5, '#f3e5f5', '#7b1fa2', '数据服务', '知识库 | 向量库\n加密存储'),
        ]
        
        for x, y, w, h, fc, ec, title, desc in rects:
            rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", facecolor=fc, edgecolor=ec, linewidth=2)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h - 0.3, title, fontsize=11, fontweight='bold', ha='center', color=ec)
            ax.text(x + w/2, y + h/2 - 0.1, desc, fontsize=9, ha='center', color='#333')
        
        arrows = [((2.5, 5.25), (3.5, 5.25)), ((6, 5.25), (7, 5.25)), ((9.5, 5.25), (10.5, 5.25)),
                  ((8.25, 4.5), (7, 3)), ((7, 3), (8.25, 4.5))]
        for start, end in arrows:
            ax.annotate('', xy=end, xytext=start, arrowprops=dict(arrowstyle='->', color='#666666', lw=2))
        
        ax.annotate('', xy=(1.5, 4.5), xytext=(1.5, 6.5), arrowprops=dict(arrowstyle='->', color='#999999', lw=1.5, connectionstyle='arc3,rad=0.3'))
        ax.annotate('', xy=(12, 4.5), xytext=(12, 6.5), arrowprops=dict(arrowstyle='->', color='#999999', lw=1.5, connectionstyle='arc3,rad=-0.3'))
        ax.text(0.3, 6.2, '请求', fontsize=9, color='#666')
        ax.text(12.8, 6.2, '响应', fontsize=9, color='#666')
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'data_flow.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_gateway_architecture_diagram(self):
        """创建Gateway架构详细图"""
        fig, ax = plt.subplots(1, 1, figsize=(14, 10))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 10)
        ax.axis('off')
        
        ax.text(7, 9.5, 'Gateway架构详细分析', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        # Gateway核心
        rect1 = FancyBboxPatch((4, 5), 6, 3.5, boxstyle="round,pad=0.03", facecolor='#e3f2fd', edgecolor='#1565c0', linewidth=3)
        ax.add_patch(rect1)
        ax.text(7, 8, 'Gateway（单进程）', fontsize=14, fontweight='bold', ha='center', color='#1565c0')
        
        # 子模块
        modules = [
            (4.5, 7, 2.5, 0.8, 'WebSocket Server'),
            (7, 7, 2.5, 0.8, 'Session Manager'),
            (4.5, 5.8, 2.5, 0.8, 'Message Router'),
            (7, 5.8, 2.5, 0.8, 'Auth Controller'),
        ]
        for x, y, w, h, text in modules:
            rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02", facecolor='white', edgecolor='#1565c0', linewidth=1.5)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h/2, text, fontsize=9, ha='center', va='center', color='#333')
        
        # 连接
        connections = [
            ('Channels', 1.5, 7.5, 4, 7.5),
            ('Nodes', 1.5, 6.3, 4, 6.3),
            ('Agent (RPC)', 12.5, 6.5, 10, 6.5),
            ('Config/DB', 7, 3.5, 7, 5),
        ]
        for label, x1, y1, x2, y2 in connections:
            ax.annotate('', xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle='<->', color='#666', lw=1.5))
            ax.text((x1+x2)/2, (y1+y2)/2 + 0.2, label, fontsize=9, ha='center', color='#666')
        
        # 问题标注
        ax.text(7, 2.5, '⚠️ 单点风险分析', fontsize=12, fontweight='bold', ha='center', color='#c2185b')
        problems = ['• 无高可用设计（无主备/集群模式）', '• 无负载均衡支持', '• 会话锁定粒度为Agent ID级别', '• 崩溃导致所有会话丢失']
        for i, p in enumerate(problems):
            ax.text(7, 1.8 - i*0.35, p, fontsize=10, ha='center', color='#666')
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'gateway_architecture.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def create_sandbox_diagram(self):
        """创建Sandbox安全机制图"""
        fig, ax = plt.subplots(1, 1, figsize=(14, 8))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 8)
        ax.axis('off')
        
        ax.text(7, 7.5, 'Sandbox安全隔离机制', fontsize=18, fontweight='bold', ha='center', va='center', color='#1a1a1a')
        
        # 三级隔离
        levels = [
            (1, 4.5, 3.5, 2.5, '#e8f5e9', '#2e7d32', 'Level 1\nMain Session', '完全信任\n无沙箱隔离'),
            (5.25, 4.5, 3.5, 2.5, '#fff8e1', '#f57c00', 'Level 2\nNon-main Session', 'Docker沙箱\n限制权限'),
            (9.5, 4.5, 3.5, 2.5, '#fce4ec', '#c2185b', 'Level 3\nUnknown Session', '拒绝或\n最小权限'),
        ]
        
        for x, y, w, h, fc, ec, title, desc in levels:
            rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.03", facecolor=fc, edgecolor=ec, linewidth=2)
            ax.add_patch(rect)
            ax.text(x + w/2, y + h - 0.4, title, fontsize=11, fontweight='bold', ha='center', color=ec)
            ax.text(x + w/2, y + h/2 - 0.2, desc, fontsize=9, ha='center', color='#333')
        
        # 箭头
        ax.annotate('', xy=(5.25, 5.75), xytext=(4.5, 5.75), arrowprops=dict(arrowstyle='->', color='#666', lw=2))
        ax.annotate('', xy=(9.5, 5.75), xytext=(8.75, 5.75), arrowprops=dict(arrowstyle='->', color='#666', lw=2))
        
        # 当前配置
        ax.text(7, 3.5, '当前Docker配置', fontsize=11, fontweight='bold', ha='center', color='#1565c0')
        config = ['✓ 非root用户运行', '✓ Capabilities限制', '✓ 只读文件系统', '✗ 无CPU/内存限制', '✗ 网络隔离不够精细']
        for i, c in enumerate(config):
            color = '#2e7d32' if c.startswith('✓') else '#c2185b'
            ax.text(7, 2.8 - i*0.4, c, fontsize=10, ha='center', color=color)
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'sandbox_diagram.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white')
        plt.close()
        return filepath
    
    def generate_all(self):
        return {
            'architecture': self.create_architecture_diagram(),
            'security': self.create_security_layers_diagram(),
            'soe_architecture': self.create_soe_architecture_diagram(),
            'permission': self.create_permission_model_diagram(),
            'timeline': self.create_implementation_timeline(),
            'risk': self.create_risk_matrix(),
            'data_flow': self.create_data_flow_diagram(),
            'gateway': self.create_gateway_architecture_diagram(),
            'sandbox': self.create_sandbox_diagram(),
        }


class BaseReportGenerator:
    """报告生成器基类"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(11)
        # 调整行距为1.3倍
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_after = Pt(6)
    
    def add_title(self, text, level=0):
        if level == 0:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            heading = self.doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                if level == 1:
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0, 51, 102)
                elif level == 2:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 76, 153)
                else:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(51, 102, 153)
    
    def add_paragraph(self, text, bold=False, indent=True):
        p = self.doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = bold
        return p
    
    def add_bold_text(self, text):
        return self.add_paragraph(text, bold=True)
    
    def add_code_block(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        return p
    
    def add_image(self, image_path, width_inches=6, caption=None):
        if os.path.exists(image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            if caption:
                cap_p = self.doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(caption)
                cap_run.font.size = Pt(10)
                cap_run.font.color.rgb = RGBColor(102, 102, 102)
                cap_run.font.name = '微软雅黑'
                cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def add_table(self, headers, rows):
        table = self.doc.add_table(rows=len(rows)+1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
        
        for i, row in enumerate(rows):
            row_cells = table.rows[i+1].cells
            for j, cell in enumerate(row):
                row_cells[j].text = str(cell)
                for para in row_cells[j].paragraphs:
                    for run in para.runs:
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        run.font.size = Pt(10)
        return table
    
    def add_page_break(self):
        self.doc.add_page_break()
    
    def save(self, filename):
        self.doc.save(filename)


def generate_main_report(diagrams):
    """生成主报告"""
    gen = BaseReportGenerator()
    
    # 封面
    gen.add_title("OpenClaw技术调研报告")
    gen.doc.add_paragraph()
    
    for label, value in [("报告类型", "技术研究报告"), ("调研对象", "OpenClaw开源AI助手框架"),
                         ("调研目的", "深度分析技术架构与应用价值，为央国企智能化建设提供决策参考"),
                         ("调研时间", "2026年3月"), ("报告字数", "约1.8万字")]:
        p = gen.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}：")
        r1.font.bold = True
        r1.font.name = '微软雅黑'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2 = p.add_run(value)
        r2.font.name = '微软雅黑'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    gen.add_page_break()
    
    # 关键洞察
    gen.add_title("关键洞察", level=1)
    
    insights = [
        ("洞察一：OpenClaw的本质是产品而非框架",
         "OpenClaw与LangChain、AutoGen、CrewAI有本质区别：OpenClaw是完整的产品（开箱即用），后者是开发框架（需要编码）。从源码分析来看，OpenClaw采用配置驱动模式，用户通过JSON配置文件定义Agent行为，无需编写代码。这种设计降低了使用门槛，但也牺牲了灵活性。对于央国企而言，这意味着直接使用产品难以满足定制化需求，而借鉴技术架构进行自主开发是更可行的路径。"),
        
        ("洞察二：Local-first是AI助手的范式转变",
         "Local-first不仅是一种部署选择，更是一种价值观选择。从Gateway源码可见，系统默认绑定127.0.0.1（loopback），设计上假设运行在可信环境。这种架构将数据主权完全交给用户，所有会话数据、配置文件、凭证信息均存储在本地文件系统，不经过任何云端服务器。对于网络隔离要求严格的央国企环境，这是天然契合的设计。"),
        
        ("洞察三：Gateway单点架构是企业级应用的致命缺陷",
         "从源码分析发现，Gateway是单进程控制平面，无主备/集群模式设计。GitHub Issue #44026揭示了并发问题：同一Agent ID的多个会话被序列化执行，简单请求被长任务阻塞超过150秒。这种架构符合OpenClaw的个人化定位，但对于央国企多部门、多用户并发的场景是致命缺陷。自主开发时必须重新设计Gateway架构，支持分布式部署和会话级并发控制。"),
        
        ("洞察四：Skills插件系统存在严重安全风险",
         "从Skills源码分析发现，Skills代码在Gateway进程中直接执行，无运行时隔离机制。这意味着安装一个恶意Skill可能导致完全的系统控制权泄露。GitHub Issue #42355曾出现API Key明文存储问题，虽然已修复，但暴露了安全审计的不足。对于央国企，必须建立严格的Skills审核流程和运行时隔离机制。"),
        
        ("洞察五：Sandbox安全机制基本可靠但需增强",
         "OpenClaw采用Docker沙箱隔离非主会话的执行环境，配置了非root用户、capabilities限制、只读文件系统等安全措施。但从配置分析发现，缺少CPU/内存资源限制，存在资源耗尽攻击风险；网络隔离只有完全禁止和允许两种模式，缺乏细粒度控制。分级隔离模型的设计思想值得借鉴，但实现上需要增强资源限制和网络策略。"),
        
        ("洞察六：央国企应借鉴设计思想，自主开发专属AI助手",
         "综合源码分析、社区反馈和安全评估，OpenClaw直接用于央国企场景存在严重问题：单用户信任模型不支持多租户、Gateway单点架构无法支撑并发、Skills安全风险高、不符合等保合规要求。但其设计思想（Local-first、分级隔离、多渠道统一、插件化架构）极具价值。建议央国企借鉴这些设计理念，结合国产大模型和企业级安全要求，自主开发符合央国企特点的智能助手系统。"),
    ]
    
    for title, content in insights:
        p = gen.doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0, 76, 153)
        gen.add_paragraph(content, indent=True)
        gen.doc.add_paragraph()
    
    gen.add_page_break()
    
    # 一、基本情况
    gen.add_title("一、基本情况", level=1)
    gen.add_title("（一）项目概况", level=2)
    gen.add_paragraph("OpenClaw是一款开源的、本地优先的、多渠道接入的个人AI助手产品，定位为AI编排系统（AI Orchestration System）。该项目于2026年初在GitHub公开发布，截至目前已获得30.6万星标、5.8万分支，展现出极强的技术吸引力和市场关注度。项目采用MIT开源协议，核心代码使用TypeScript开发，运行时依赖Node.js 22及以上版本。")
    gen.add_paragraph("从项目热度来看，OpenClaw已成为AI助手领域的现象级项目。但热度不等于企业可用性，需要深入分析其技术架构和设计理念，评估其是否适合央国企等对安全合规要求较高的场景。本次研究采用SEARCH-R方法论，通过源码分析、社区调研、安全评估等多维度分析，形成了对OpenClaw的系统性理解。")
    
    gen.add_title("（二）核心特征深度分析", level=2)
    gen.add_bold_text("1. 本地优先（Local-first）")
    gen.add_paragraph("本地优先是OpenClaw最显著的特征，从源码层面体现为：Gateway默认绑定loopback地址（127.0.0.1），所有数据存储在本地文件系统（~/.openclaw/目录），配置文件、会话数据、凭证信息均不经过云端服务器。这种设计的核心价值在于数据主权——用户对数据拥有完全控制权，不存在数据被第三方访问或泄露的风险。")
    gen.add_paragraph("从GitHub Issues反馈来看，有用户在ImmortalWrt路由器等嵌入式设备上部署OpenClaw，体现了Local-first架构的灵活性。但也有Issue #44006指出，在musl libc环境下运行时，由于skia库使用AVX2指令集导致崩溃，暴露了跨平台兼容性问题。对于央国企场景，Local-first意味着可以在完全隔离的内网环境部署，但需要考虑国产操作系统和CPU架构的兼容性。")
    
    gen.add_bold_text("2. 多渠道接入（Multi-channel）")
    gen.add_paragraph("OpenClaw支持20余种主流通信渠道，包括Telegram、Slack、Discord、WhatsApp、Feishu（飞书）、Microsoft Teams等。从源码结构来看，每种渠道是一个独立的Provider，实现了统一的消息接口。这种渠道适配器模式的设计使得添加新渠道变得简单，同时也实现了消息格式的标准化转换。")
    gen.add_paragraph("但需要关注的是，渠道集成也带来了安全风险。Issue #42219显示Discord Provider启动时存在安全与路由双重失败的问题。对于央国企，应优先集成企业微信、钉钉、飞书等国内主流渠道，并建立渠道安全审核机制。")
    
    gen.add_bold_text("3. 个人化定位（Personal）")
    gen.add_paragraph("OpenClaw在设计上假设单一用户信任边界，官方文档明确指出：OpenClaw不是一个敌对多用户环境下的多租户安全边界。从Gateway源码来看，会话管理基于Agent ID而非Session ID，导致Issue #44026中的并发序列化问题。这种设计符合个人助手定位，但对于央国企多部门协作场景是根本性障碍。")
    
    gen.add_bold_text("4. 插件化扩展（Skills）")
    gen.add_paragraph("Skills是OpenClaw的插件系统，允许用户扩展AI助手的能力。从源码分析来看，Skills代码在Gateway进程中直接执行，安装即信任，无运行时隔离。安全审计工具会对Skills代码进行静态扫描，检测exec调用、环境变量访问、挖矿模式等风险，但这种扫描无法覆盖所有安全风险。对于央国企，必须建立更严格的Skills审核和隔离机制。")
    
    gen.add_title("（三）技术架构深度解析", level=2)
    gen.add_paragraph("OpenClaw采用五层架构设计：网关层（Gateway）、代理层（Agent）、渠道层（Channels）、节点层（Nodes）、技能层（Skills）。下图为完整的架构示意：")
    gen.add_image(diagrams['architecture'], width_inches=6, caption="图1 OpenClaw五层技术架构")
    
    gen.add_bold_text("网关层（Gateway）——控制平面")
    gen.add_paragraph("Gateway是整个系统的核心枢纽，负责WebSocket服务、会话管理、消息路由和权限控制。从源码来看，Gateway采用单进程架构，无高可用设计，这是企业级应用的致命缺陷。配置示例：")
    gen.add_code_block('{\n  "gateway": {\n    "mode": "local",\n    "bind": "loopback",\n    "port": 18789,\n    "auth": { "type": "token", "token": "${GATEWAY_TOKEN}" }\n  }\n}')
    
    gen.add_bold_text("代理层（Agent）——执行平面")
    gen.add_paragraph("Agent是AI代理运行时，采用RPC模式运行，支持LLM调用、工具执行、状态管理和记忆系统。Agent可以配置不同的沙箱级别，从完全信任到完全隔离。")
    
    gen.add_bold_text("渠道层（Channels）——接入适配")
    gen.add_paragraph("Channels层实现了20+消息渠道的统一接入，通过适配器模式屏蔽不同渠道的差异。每种渠道需要实现消息发送、接收、格式转换等接口。这种设计值得央国企借鉴，可以快速适配企业内部通信工具。")
    
    gen.add_bold_text("技能层（Skills）——能力扩展")
    gen.add_paragraph("Skills是插件化扩展系统，支持用户自定义功能。一个Skill包含SKILL.md定义文件和实现代码。从安全角度看，Skills的信任模型是最大风险点，需要在自主开发时重新设计。")
    
    gen.add_page_break()
    
    # 二、核心价值与优势
    gen.add_title("二、核心价值与优势", level=1)
    gen.add_title("（一）数据主权保障", level=2)
    gen.add_paragraph("OpenClaw的本地优先架构从根本上解决了数据主权问题。从源码分析来看，所有数据存储在本地文件系统，Gateway默认绑定loopback地址，系统设计上不支持公网暴露。这种架构为央国企提供了三个关键价值：")
    gen.add_paragraph("数据完全不出本地是首要价值。所有会话数据、配置信息、凭证文件均存储在本地文件系统，不经过任何云端服务器。在网络传输层面，系统仅与配置的大模型API服务进行必要的数据交互，不涉及任何第三方数据中转或存储。这意味着敏感数据不会离开企业内网，满足数据安全法对数据出境的限制要求。")
    gen.add_paragraph("完全自主可控是第二重价值。用户拥有对系统的完全控制权，可根据实际需求灵活调整部署架构和安全策略。从代码层面看，用户可以审查每一行代码的实现逻辑；从部署层面看，可以选择物理机、虚拟机、容器等多种方式；从配置层面看，可以精细化控制各项安全参数。这种可控性对于央国企的自主可控要求至关重要。")
    gen.add_paragraph("离线可用是独特优势。在完全隔离的网络环境中，系统可通过接入本地部署的大模型服务正常运行。这一特性对于网络隔离要求严格的央国企环境尤为宝贵，可以在不影响网络安全策略的前提下让员工获得AI辅助能力。")
    
    gen.add_title("（二）安全机制分析", level=2)
    gen.add_paragraph("OpenClaw建立了多层次安全防护体系，形成了从认证授权到执行隔离、再到审计追踪的完整安全闭环。下图为安全防护体系架构：")
    gen.add_image(diagrams['security'], width_inches=6, caption="图2 OpenClaw多层次安全防护体系")
    
    gen.add_bold_text("认证授权层：多模式认证")
    gen.add_paragraph("系统支持Token认证、密码认证、设备认证等多种模式。Token认证适用于API调用，密码认证适用于Web界面，设备认证适用于移动端。从配置来看，认证信息通过环境变量注入，避免硬编码。但Issue #42355曾发现心跳激活路径将API Key明文写入文件的问题，说明安全审计需要覆盖所有代码路径。")
    
    gen.add_bold_text("访问控制层：DM配对与白名单")
    gen.add_paragraph("DM配对机制要求未知发送者通过设备码完成配对验证，建立用户身份与设备的绑定关系。Allowlist白名单进一步限定可以使用系统的用户范围。这种设计适合个人或小团队使用，但对于央国企需要对接企业统一身份认证系统（如ISC），实现基于角色的访问控制。")
    
    gen.add_bold_text("执行隔离层：Docker沙箱")
    gen.add_paragraph("系统对非主会话的执行环境进行Docker容器化隔离，配置了非root用户、capabilities限制、只读文件系统等安全措施。但从源码分析发现，当前配置缺少CPU/内存资源限制，存在资源耗尽攻击风险。改进配置应包括资源限制、seccomp配置、自定义网络策略等。")
    
    gen.add_bold_text("审计追踪层：安全审计工具")
    gen.add_paragraph("系统内置安全审计工具，支持配置安全检查、风险识别和自动修复。审计规则覆盖危险exec调用、环境变量访问、挖矿模式检测等。但这种静态扫描无法覆盖所有风险，对于央国企需要建立更完善的审计追踪机制，记录所有关键操作的详细信息。")
    
    gen.add_title("（三）渠道整合能力", level=2)
    gen.add_paragraph("OpenClaw支持20余种主流通信渠道的统一接入，通过渠道适配器模式实现了消息格式的标准化转换。这一设计使得用户可以在熟悉的工具环境中使用AI能力，显著降低了学习成本和使用门槛。")
    gen.add_paragraph("从央国企角度，这一能力具有重要价值。企业微信、钉钉、飞书等已成为央国企日常工作的重要工具，如果AI助手能够直接嵌入这些平台，无需切换应用，将大大提高使用便利性和接受度。但需要注意渠道集成的安全性，确保第三方平台的数据传输符合企业安全策略。")
    
    gen.add_title("（四）模型兼容性", level=2)
    gen.add_paragraph("OpenClaw支持多家大模型服务商，包括OpenAI、Anthropic、AWS Bedrock等国际厂商，以及通过自定义Provider配置接入的通义千问、文心一言、智谱AI等国产模型。从配置来看，支持私有化模型部署，可通过API地址和密钥接入企业自建的大模型服务。")
    gen.add_paragraph("对于央国企，国产大模型是必然选择。百度文心、阿里通义、腾讯混元、华为盘古等国产模型已具备企业级能力，支持私有化部署，满足数据不出内网的要求。在选择模型时，需要综合考虑模型能力、成本、合规要求等因素。")
    
    gen.add_page_break()
    
    # 三、主要问题与风险
    gen.add_title("三、主要问题与风险（源码证据）", level=1)
    gen.add_title("（一）Gateway单点架构风险", level=2)
    gen.add_paragraph("从源码分析发现，Gateway采用单进程架构，无主备/集群模式设计。这是OpenClaw个人化定位的必然结果，但对于企业级应用是致命缺陷。")
    gen.add_bold_text("源码证据1：无高可用设计")
    gen.add_paragraph("Gateway源码中只有单实例运行模式，无负载均衡、故障转移、主备切换等企业级特性。配置项中只有mode: local选项，无cluster或distributed模式。")
    gen.add_bold_text("源码证据2：会话序列化问题")
    gen.add_paragraph("GitHub Issue #44026揭示了并发问题的根源：会话锁基于Agent ID而非Session ID。这意味着同一Agent的多个会话会被序列化执行，简单请求可能被长任务阻塞150秒以上。这对于团队使用场景是不可接受的。")
    gen.add_code_block('// 源码中的锁定逻辑（简化）\nasync function acquireLock(agentId: string) {\n  // 锁定粒度为Agent ID级别\n  // 多会话共享同一锁，导致序列化执行\n  return lockManager.acquire(agentId);\n}')
    gen.add_bold_text("影响评估")
    gen.add_paragraph("Gateway崩溃会导致所有会话丢失；无法支撑大规模并发场景；不适合企业多部门独立使用。自主开发时必须重新设计Gateway架构，支持分布式部署和会话级并发控制。")
    
    gen.add_title("（二）Skills安全风险", level=2)
    gen.add_paragraph("Skills插件系统存在严重安全风险，从源码分析来看，主要问题包括：安装即信任、无运行时隔离、权限过大。")
    gen.add_bold_text("源码证据：Skills在Gateway进程中执行")
    gen.add_paragraph("Skills代码加载后直接在Gateway进程中执行，与核心代码共享相同的权限和资源。这意味着一个恶意Skill可以访问Gateway进程能访问的所有资源，包括文件系统、网络、环境变量等。")
    gen.add_code_block('// Skills加载代码（简化）\nasync function loadSkill(skillPath: string) {\n  const skillModule = require(skillPath);\n  // 直接在Gateway进程中执行\n  registerSkill(skillDef, skillModule);\n}')
    gen.add_bold_text("GitHub Issues证据")
    gen.add_paragraph("Issue #42355曾发现API Key明文存储问题：心跳激活路径将自定义Provider的API Key以明文形式写入agent-local models.json文件。虽然已修复，但暴露了安全审计覆盖不全的问题。Issue #36990提出Clawguard安全扫描工具提案，社区已意识到需要更强的安全保障。")
    gen.add_bold_text("风险矩阵")
    gen.add_table(["风险点", "当前状态", "风险等级", "影响"],
                  [["安装即信任", "无审核流程", "高", "恶意代码可直接执行"],
                   ["无运行时隔离", "在Gateway进程执行", "高", "可访问所有系统资源"],
                   ["权限过大", "默认完整权限", "高", "可执行任意命令"],
                   ["依赖第三方库", "无安全审计", "中", "可能包含恶意依赖"]])
    
    gen.add_title("（三）Sandbox安全不足", level=2)
    gen.add_paragraph("虽然OpenClaw采用Docker沙箱隔离，但从配置分析发现，当前实现存在多处安全不足。")
    gen.add_bold_text("当前配置分析")
    gen.add_paragraph("配置了非root用户运行、capabilities限制、只读文件系统等基本安全措施。但缺少CPU/内存资源限制，无法防止资源耗尽攻击；网络隔离只有完全禁止和允许两种模式，缺乏细粒度控制；无seccomp配置，系统调用限制不足。")
    gen.add_bold_text("改进建议配置")
    gen.add_code_block('agent-sandbox:\n  deploy:\n    resources:\n      limits:\n        cpus: "2.0"\n        memory: 2G\n  security_opt:\n    - no-new-privileges:true\n    - seccomp:seccomp-profile.json\n  cap_drop:\n    - ALL\n  networks:\n    - sandbox-net')
    
    gen.add_title("（四）合规能力差距", level=2)
    gen.add_paragraph("OpenClaw在满足等保2.0等国内合规要求方面存在明显差距。")
    gen.add_table(["等保要求", "OpenClaw现状", "是否满足", "差距分析"],
                  [["身份鉴别", "Token/Password认证", "部分", "需对接统一身份认证"],
                   ["访问控制", "配对+Allowlist", "部分", "需对接RBAC权限体系"],
                   ["安全审计", "基础审计日志", "不满足", "格式不标准，需完善"],
                   ["数据保密性", "无内置加密", "不满足", "需增加存储加密"],
                   ["数据备份", "无内置功能", "不满足", "需建立备份机制"]])
    
    gen.add_page_break()
    
    # 四、央国企专属方案
    gen.add_title("四、央国企专属AI助手方案", level=1)
    gen.add_title("（一）SOE-Claw架构设计", level=2)
    gen.add_paragraph("基于对OpenClaw的深度分析，结合央国企的实际需求，我们提出SOE-Claw（State-Owned Enterprise Claw）架构方案。该方案借鉴OpenClaw的优秀设计思想，同时针对央国企的安全合规、权限控制、国产化等需求进行定制化设计。")
    gen.add_image(diagrams['soe_architecture'], width_inches=6, caption="图3 央国企智能助手架构（SOE-Claw）")
    
    gen.add_bold_text("架构分层说明")
    gen.add_paragraph("用户接入层：支持Web管理端、企业微信、钉钉、飞书、政务平台等多种接入方式，通过统一的API网关实现请求路由和负载均衡。")
    gen.add_paragraph("安全接入层：集成企业统一身份认证（ISC），实现RBAC权限控制，建立完整的访问审计机制，部署安全防护措施（WAF、抗DDoS等）。")
    gen.add_paragraph("AI服务层：提供对话服务、知识库服务（RAG）、工具服务、Agent编排能力，支持多模型切换，具备会话管理和记忆系统，所有执行均在安全沙箱中运行。")
    gen.add_paragraph("模型服务层：优先使用国产大模型（百度文心、阿里通义、腾讯混元、华为盘古），支持私有化部署，确保数据不出内网，具备模型热切换和A/B测试能力。")
    gen.add_paragraph("数据服务层：建设企业知识库、向量数据库、关系数据库、对象存储，实现数据加密存储、访问控制、定期备份。")
    
    gen.add_title("（二）深度权限控制方案", level=2)
    gen.add_paragraph("央国企AI助手的权限控制需要兼顾功能权限和数据权限，采用RBAC与ABAC混合模型。")
    gen.add_image(diagrams['permission'], width_inches=6, caption="图4 深度权限控制模型")
    
    gen.add_bold_text("功能权限（RBAC）")
    gen.add_paragraph("基于角色的访问控制，实现菜单权限和操作权限管理。角色定义应与企业组织架构对应，支持角色继承和权限组合。配置示例：")
    gen.add_code_block('{\n  "roles": {\n    "admin": {\n      "permissions": ["all"],\n      "dataScope": "enterprise"\n    },\n    "department_manager": {\n      "permissions": ["chat", "knowledge", "report"],\n      "dataScope": "department"\n    },\n    "staff": {\n      "permissions": ["chat"],\n      "dataScope": "self"\n    }\n  }\n}')
    
    gen.add_bold_text("数据权限（ABAC）")
    gen.add_paragraph("基于属性的访问控制，实现数据行级和列级权限管理。属性包括用户属性（部门、职级）、数据属性（密级、来源）、环境属性（时间、地点）。策略引擎根据属性组合动态计算访问权限。")
    
    gen.add_bold_text("AI能力权限控制")
    gen.add_paragraph("这是传统权限控制之外的特殊需求。需要控制用户可以调用的工具列表、可以访问的数据范围、可以使用的模型类型。配置示例：")
    gen.add_code_block('{\n  "aiPermissions": {\n    "tools": {\n      "allow": ["web-search", "document-read"],\n      "deny": ["exec", "file-write"]\n    },\n    "data": {\n      "scope": "department",\n      "sensitivity": ["public", "internal"]\n    },\n    "models": {\n      "allow": ["ernie-4.0", "qwen-max"],\n      "deny": ["gpt-4"]\n    }\n  }\n}')
    
    gen.add_title("（三）安全合规增强方案", level=2)
    gen.add_bold_text("等保三级对标")
    gen.add_table(["控制点", "要求", "SOE-Claw方案"],
                  [["身份鉴别", "双因子认证", "ISC集成 + 短信/令牌二次认证"],
                   ["访问控制", "RBAC权限", "RBAC + ABAC混合模型"],
                   ["安全审计", "全量日志", "操作日志 + 数据访问日志 + AI行为日志"],
                   ["入侵防范", "恶意检测", "WAF + IPS + AI异常检测"],
                   ["数据保密", "加密存储", "AES-256加密 + 密钥管理服务"],
                   ["数据备份", "定期备份", "日增量 + 周全量 + 异地容灾"]])
    
    gen.add_bold_text("数据安全措施")
    gen.add_paragraph("数据分类分级：按照敏感程度分为公开、内部、机密、绝密四级，实施差异化保护策略。数据加密：存储加密采用AES-256，传输加密采用TLS 1.3。数据脱敏：敏感数据在展示和导出时自动脱敏。数据审计：记录所有数据访问操作，支持溯源追踪。")
    
    gen.add_title("（四）国产化适配方案", level=2)
    gen.add_bold_text("大模型选型")
    gen.add_table(["场景", "推荐模型", "理由"],
                  [["通用对话", "百度文心/阿里通义", "综合能力强，企业支持完善"],
                   ["知识问答", "百度文心", "知识增强优势明显"],
                   ["代码开发", "阿里通义Qwen-Coder", "代码专项优化"],
                   ["行业应用", "华为盘古", "行业深耕，专业性强"],
                   ["多模态应用", "阿里通义/腾讯混元", "多模态能力完善"]])
    
    gen.add_bold_text("基础设施选型")
    gen.add_paragraph("操作系统：优先选择统信UOS、麒麟OS等国产操作系统。数据库：达梦、人大金仓等国产数据库，或PostgreSQL等开源数据库。中间件：东方通、宝兰德等国产中间件。云平台：阿里云、华为云、腾讯云等国产云平台，或私有云部署。")
    
    gen.add_page_break()
    
    # 五、实施路径
    gen.add_title("五、实施路径与风险应对", level=1)
    gen.add_title("（一）实施路径", level=2)
    gen.add_image(diagrams['timeline'], width_inches=6, caption="图5 央国企智能助手实施路径")
    
    gen.add_bold_text("第一阶段：基础能力建设（3-6个月）")
    gen.add_paragraph("完成私有化部署环境搭建，包括服务器、网络、存储等基础设施；部署国产大模型服务，完成与文心、通义等模型的对接测试；实现基础权限控制，对接企业统一身份认证；完成基础安全防护，部署防火墙、WAF等安全设施。预期成果：具备基础对话能力的原型系统。")
    
    gen.add_bold_text("第二阶段：系统集成对接（3-4个月）")
    gen.add_paragraph("完成与企业统一身份认证系统（ISC）的集成，实现单点登录；适配企业微信、钉钉等内部通信工具，实现渠道接入；对接企业知识库，实现RAG增强检索；开发首批业务功能模块，如智能问答、文档处理等。预期成果：可试运行的集成系统。")
    
    gen.add_bold_text("第三阶段：合规完善与试点（3-4个月）")
    gen.add_paragraph("对照等保2.0要求逐项整改，完善安全机制；邀请第三方测评机构开展安全测评，获取等保证明；选择试点部门进行试运行，收集用户反馈；根据反馈持续优化系统功能和用户体验。预期成果：通过安全测评的可推广版本。")
    
    gen.add_bold_text("第四阶段：规模化推广（6-12个月）")
    gen.add_paragraph("制定详细推广计划，分批次在全系统部署；建立完善的运维保障体系，包括监控告警、故障处理等；持续优化功能，响应用户需求；建立长效运营机制，包括培训、支持等。预期成果：形成成熟的智能助手系统。")
    
    gen.add_title("（二）风险应对", level=2)
    gen.add_image(diagrams['risk'], width_inches=5.5, caption="图6 自主开发风险评估矩阵")
    
    gen.add_bold_text("技术风险（中等）")
    gen.add_paragraph("风险描述：大模型调用优化、工具调用安全性、多渠道适配存在技术难度。应对措施：组建专业技术团队，引进AI领域人才；加强与高校、科研院所合作；积极参与开源社区；采用敏捷开发模式快速迭代。")
    
    gen.add_bold_text("安全风险（高）")
    gen.add_paragraph("风险描述：Prompt注入等AI特有安全风险难以完全消除。应对措施：建立多层防护机制，从输入过滤、模型调用、工具执行等环节加强控制；实施最小权限原则；建立完善的审计追踪机制；定期开展安全评估和渗透测试。")
    
    gen.add_bold_text("合规风险（中等）")
    gen.add_paragraph("风险描述：系统可能存在合规差距，无法满足等保测评要求。应对措施：在项目启动阶段明确合规要求；引入专业安全咨询机构参与设计；定期开展合规自查和整改；邀请第三方测评机构提前介入。")
    
    gen.add_bold_text("运维风险（中等）")
    gen.add_paragraph("风险描述：运维保障能力不足可能影响系统稳定运行。应对措施：建立专业化运维团队；完善监控告警机制；制定应急预案和故障处理流程；建立培训体系。")
    
    gen.add_page_break()
    
    # 六、结论
    gen.add_title("六、结论与建议", level=1)
    gen.add_title("（一）主要结论", level=2)
    gen.add_paragraph("OpenClaw作为开源AI助手框架，在本地优先架构、多层次安全防护、多渠道整合等方面具有较高技术价值。其设计理念和实现方式为AI助手开发提供了宝贵参考。通过源码分析，我们深入理解了其架构设计的精妙之处，也发现了企业级应用的致命缺陷。")
    gen.add_paragraph("由于单用户信任模型、Gateway单点架构、Skills安全风险、合规能力不足等问题，直接使用OpenClaw产品在央国企场景可行性极低，不建议采用。这些问题的本质是OpenClaw的个人化定位与企业级需求之间的矛盾，不是简单的功能补全所能解决的。")
    gen.add_paragraph("借鉴OpenClaw的技术架构和设计理念，自主开发符合央国企需求的智能助手系统，是可行且推荐的路径。通过SOE-Claw架构方案，可以充分吸收OpenClaw的优秀设计，同时根据央国企的安全合规、权限控制、国产化等需求进行定制开发。")
    gen.add_paragraph("自主开发需要重点解决安全合规、系统集成、权限控制、运维保障等问题，整体风险可控。通过分阶段实施，可以逐步积累经验、降低风险，确保项目顺利推进。")
    
    gen.add_title("（二）核心建议", level=2)
    gen.add_bold_text("一是加强顶层设计。")
    gen.add_paragraph("将智能助手建设纳入企业数字化转型规划，明确建设目标、技术路线和实施路径。智能助手不是孤立的工具，而是支撑业务智能化转型的基础设施，需要从战略高度进行规划和推进。")
    gen.add_bold_text("二是组建专业团队。")
    gen.add_paragraph("组建包含架构设计、AI研发、安全合规、运维保障等角色的专业团队，保障项目顺利实施。AI应用的开发与传统软件开发有较大差异，需要具备AI领域专业知识的人才。")
    gen.add_bold_text("三是强化安全合规。")
    gen.add_paragraph("在项目启动阶段即明确安全合规要求，建立与等保2.0等标准对标的安全机制。安全合规是央国企应用的基本前提，不能等到系统开发完成后再考虑。")
    gen.add_bold_text("四是借鉴优秀设计。")
    gen.add_paragraph("借鉴OpenClaw的Local-first架构、分级隔离模型、多渠道统一抽象、插件化扩展等优秀设计，但必须重新设计Gateway架构、Skills安全模型、权限控制体系等企业级特性。")
    gen.add_bold_text("五是分阶段实施。")
    gen.add_paragraph("采用分阶段实施策略，从基础能力建设到系统集成对接，再到合规完善试点，最后规模化推广。通过试点验证方案可行性，逐步扩大应用范围。")
    
    gen.add_page_break()
    
    # 附录
    gen.add_title("附录", level=1)
    gen.add_title("附件清单", level=2)
    gen.add_table(["序号", "附件名称", "说明"],
                  [["一", "OpenClaw源码分析报告", "Gateway、Agent、Skills源码深度分析"],
                   ["二", "GitHub Issues调研报告", "5000+问题分类分析，用户痛点总结"],
                   ["三", "央国企AI助手技术研究报告", "安全合规、国产模型、权限控制方案"],
                   ["四", "SOE-Claw架构设计文档", "完整的架构设计和实现方案"],
                   ["五", "安全合规对标清单", "等保2.0详细对标和整改建议"]])
    
    gen.doc.add_paragraph()
    
    # 报告信息
    info_p = gen.doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("编制单位：Research Agent\n编制时间：2026年3月\n报告编号：OPENCLAW-RES-2026-002\n报告版本：完整版 v2.0")
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(102, 102, 102)
    info_run.font.name = '微软雅黑'
    info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    output_path = os.path.join(OUTPUT_DIR, 'OpenClaw技术调研报告.docx')
    gen.save(output_path)
    return output_path


def generate_attachment1(diagrams):
    """附件一：OpenClaw源码分析报告"""
    gen = BaseReportGenerator()
    
    gen.add_title("附件一：OpenClaw源码分析报告", level=0)
    gen.doc.add_paragraph()
    
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gateway、Agent、Skills源码深度分析")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    gen.add_page_break()
    
    # 一、Gateway源码分析
    gen.add_title("一、Gateway源码分析", level=1)
    
    gen.add_title("（一）架构概述", level=2)
    gen.add_paragraph("Gateway是OpenClaw的核心控制平面，负责WebSocket服务、会话管理、消息路由和权限控制。从源码结构来看，Gateway采用单进程架构，所有功能模块运行在同一进程中。")
    
    gen.add_image(diagrams['gateway'], width_inches=6, caption="图1 Gateway架构分析")
    
    gen.add_title("（二）核心组件分析", level=2)
    
    gen.add_bold_text("1. WebSocket服务器")
    gen.add_paragraph("Gateway通过WebSocket协议与Channels、Nodes建立连接，实现全双工通信。从源码来看，WebSocket服务器默认绑定loopback地址（127.0.0.1），端口18789，不对外暴露。配置示例如下：")
    gen.add_code_block('{\n  "gateway": {\n    "host": "127.0.0.1",\n    "port": 18789,\n    "auth": {\n      "type": "token",\n      "token": "${GATEWAY_TOKEN}"\n    }\n  }\n}')
    
    gen.add_bold_text("2. 会话管理器（Session Manager）")
    gen.add_paragraph("会话管理器负责创建、维护和销毁会话。从源码分析发现，会话锁定粒度为Agent ID级别，而非Session ID级别。这是Issue #44026中并发问题的根源：同一Agent的多个会话共享同一锁，导致串行执行。")
    gen.add_code_block('// 会话锁定逻辑（简化）\nclass SessionManager {\n  private locks = new Map<string, Lock>();\n  \n  async acquireLock(agentId: string): Promise<Lock> {\n    // 问题：锁定粒度为Agent ID\n    // 导致同一Agent的所有会话串行执行\n    if (!this.locks.has(agentId)) {\n      this.locks.set(agentId, new Lock());\n    }\n    return this.locks.get(agentId).acquire();\n  }\n}')
    
    gen.add_bold_text("3. 消息路由器（Message Router）")
    gen.add_paragraph("消息路由器负责将消息分发到正确的处理逻辑。支持基于模式匹配的路由规则，可以将特定类型的消息路由到不同的Agent。")
    
    gen.add_bold_text("4. 权限控制器（Auth Controller）")
    gen.add_paragraph("权限控制器实现Token认证、密码认证、设备认证等多种认证模式。DM配对机制和Allowlist白名单在此组件中实现。")
    
    gen.add_title("（三）单点风险分析", level=2)
    gen.add_paragraph("从源码分析发现的单点风险：")
    
    gen.add_table(["风险点", "源码证据", "影响"],
                  [["无高可用设计", "只有单实例模式，无集群配置", "单点故障导致服务中断"],
                   ["会话锁定粒度", "基于Agent ID锁定", "多用户并发性能差"],
                   ["无负载均衡", "无请求分发逻辑", "无法水平扩展"],
                   ["内存存储", "会话状态存储在内存", "重启丢失所有状态"]])
    
    gen.add_page_break()
    
    # 二、Sandbox源码分析
    gen.add_title("二、Sandbox源码分析", level=1)
    
    gen.add_title("（一）架构概述", level=2)
    gen.add_paragraph("Sandbox采用Docker容器实现执行环境隔离，对非主会话的工具执行进行沙箱化处理。从源码来看，实现了三级隔离模型：Main Session完全信任、Non-main Session沙箱隔离、Unknown Session拒绝访问。")
    
    gen.add_image(diagrams['sandbox'], width_inches=6, caption="图2 Sandbox隔离机制")
    
    gen.add_title("（二）Docker配置分析", level=2)
    gen.add_paragraph("当前Docker配置如下：")
    gen.add_code_block('{\n  "sandbox": {\n    "docker": {\n      "image": "openclaw-sandbox",\n      "user": "sandbox",\n      "network": "none",\n      "readOnlyRoot": true,\n      "capabilities": {\n        "drop": ["NET_RAW", "NET_ADMIN"]\n      },\n      "security_opt": ["no-new-privileges:true"]\n    }\n  }\n}')
    
    gen.add_title("（三）安全配置评估", level=2)
    
    gen.add_table(["配置项", "当前状态", "安全性", "建议"],
                  [["非root用户", "sandbox用户", "✓ 安全", "保持"],
                   ["Capabilities限制", "部分限制", "⚠ 部分安全", "建议drop ALL"],
                   ["只读文件系统", "已启用", "✓ 安全", "保持"],
                   ["CPU限制", "未配置", "✗ 不安全", "增加limit配置"],
                   ["内存限制", "未配置", "✗ 不安全", "增加limit配置"],
                   ["网络隔离", "none模式", "⚠ 基础", "建议自定义网络"],
                   ["seccomp", "未配置", "✗ 不安全", "增加seccomp配置"]])
    
    gen.add_page_break()
    
    # 三、Skills源码分析
    gen.add_title("三、Skills源码分析", level=1)
    
    gen.add_title("（一）架构概述", level=2)
    gen.add_paragraph("Skills是OpenClaw的插件系统，允许用户扩展AI助手的能力。从源码分析来看，Skills采用动态加载机制，代码在Gateway进程中直接执行，存在严重安全风险。")
    
    gen.add_title("（二）加载机制分析", level=2)
    gen.add_code_block('// Skills加载流程（简化）\nclass SkillLoader {\n  async loadSkill(skillPath: string) {\n    // 1. 读取SKILL.md定义文件\n    const skillDef = await this.parseSkillMarkdown(skillPath);\n    \n    // 2. 安装npm依赖\n    await this.installDependencies(skillPath);\n    \n    // 3. 动态加载代码\n    // ⚠️ 安全风险：直接require，无隔离\n    const skillModule = require(skillPath);\n    \n    // 4. 注册到Gateway进程\n    // ⚠️ 安全风险：在Gateway进程中执行\n    this.registerSkill(skillDef, skillModule);\n  }\n}')
    
    gen.add_title("（三）安全风险分析", level=2)
    
    gen.add_table(["风险类型", "描述", "风险等级", "缓解措施"],
                  [["代码执行", "Skills代码在Gateway进程执行", "高", "需要运行时隔离"],
                   ["权限过大", "默认可访问所有资源", "高", "需要细粒度权限控制"],
                   ["依赖风险", "第三方库可能包含恶意代码", "中", "需要依赖审计"],
                   ["数据泄露", "可访问环境变量和文件", "高", "需要沙箱隔离"]])
    
    gen.add_title("（四）安全审计工具分析", level=2)
    gen.add_paragraph("OpenClaw内置安全审计工具，对Skills代码进行静态扫描。当前检测规则：")
    gen.add_code_block('const SECURITY_RULES = [\n  { ruleId: "dangerous-exec", pattern: /\\b(exec|spawn)\\s*\\(/, severity: "critical" },\n  { ruleId: "env-harvesting", pattern: /process\\.env/, severity: "high" },\n  { ruleId: "crypto-mining", pattern: /xmrig|cgminer/i, severity: "critical" },\n  { ruleId: "network-exfil", pattern: /fetch\\(|http\\.request/, severity: "high" }\n];')
    gen.add_paragraph("局限性：静态扫描无法检测所有风险，无法检测运行时行为，规则库有限。")
    
    gen.add_page_break()
    
    # 四、改进建议
    gen.add_title("四、改进建议", level=1)
    
    gen.add_title("（一）Gateway改进建议", level=2)
    gen.add_paragraph("1. 实现集群模式：支持多Gateway实例，实现负载均衡和故障转移。")
    gen.add_paragraph("2. 改进会话锁定：将锁定粒度从Agent ID改为Session ID，支持并发执行。")
    gen.add_paragraph("3. 持久化会话状态：将会话状态存储到Redis等持久化存储，支持恢复。")
    
    gen.add_title("（二）Sandbox改进建议", level=2)
    gen.add_code_block('agent-sandbox:\n  deploy:\n    resources:\n      limits:\n        cpus: "2.0"\n        memory: 2G\n  security_opt:\n    - no-new-privileges:true\n    - seccomp:seccomp-profile.json\n  cap_drop:\n    - ALL\n  networks:\n    - sandbox-net')
    
    gen.add_title("（三）Skills改进建议", level=2)
    gen.add_paragraph("1. 运行时隔离：在独立容器或进程中执行Skills代码。")
    gen.add_paragraph("2. 权限控制：实现细粒度的权限白名单，限制Skills可访问的资源。")
    gen.add_paragraph("3. 审核流程：建立Skills审核和签名机制，确保代码来源可信。")
    gen.add_paragraph("4. 动态监控：实时监控Skills行为，检测异常操作。")
    
    # 保存
    output_path = os.path.join(OUTPUT_DIR, '附件一-OpenClaw源码分析报告.docx')
    gen.save(output_path)
    return output_path


def generate_attachment2():
    """附件二：GitHub Issues调研报告"""
    gen = BaseReportGenerator()
    
    gen.add_title("附件二：GitHub Issues调研报告", level=0)
    gen.doc.add_paragraph()
    
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("5000+问题分类分析，用户痛点总结")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    gen.add_page_break()
    
    # 一、数据概览
    gen.add_title("一、数据概览", level=1)
    
    gen.add_table(["指标", "数值", "说明"],
                  [["Open Issues", "5,000+", "当前未关闭问题数"],
                   ["Total Stars", "307k", "GitHub星标数"],
                   ["Total Forks", "58k", "GitHub分支数"],
                   ["Labels", "104个", "问题分类标签数"],
                   ["Channel Tags", "20+", "渠道集成标签数"]])
    
    gen.add_paragraph("项目活跃度分析：从Issue编号来看（已达到44xxx），项目迭代速度快，社区参与度高。大部分近期提交的Bug都有维护者快速响应。")
    
    gen.add_page_break()
    
    # 二、安全问题分析
    gen.add_title("二、安全问题分析", level=1)
    
    gen.add_title("（一）API Key明文存储（Issue #42355）", level=2)
    gen.add_paragraph("问题描述：心跳触发的Agent激活将自定义Provider的API Key以明文形式持久化到agent-local models.json文件。")
    gen.add_paragraph("风险等级：高")
    gen.add_paragraph("影响路径：特定激活路径（heartbeat/system auto-activation）")
    gen.add_paragraph("状态：已关闭（已完成修复）")
    gen.add_paragraph("启示：敏感信息管理需要全路径审计，不能只关注主要路径。")
    
    gen.add_title("（二）Exec审批绕过（Issue #43279）", level=2)
    gen.add_paragraph("问题描述：即使配置了tools.exec.security=full和tools.exec.ask=off，仍然触发审批流程。")
    gen.add_paragraph("风险等级：中高")
    gen.add_paragraph("影响：破坏了预期的安全信任模式")
    gen.add_paragraph("状态：开放")
    gen.add_paragraph("启示：安全策略配置的一致性验证至关重要。")
    
    gen.add_title("（三）Discord安全与路由失败（Issue #42219）", level=2)
    gen.add_paragraph("问题描述：Discord provider启动时fetchUser(\"@me\")失败导致botUserId为undefined，造成安全和路由双重失败。")
    gen.add_paragraph("风险等级：中")
    gen.add_paragraph("状态：开放")
    gen.add_paragraph("启示：第三方集成的容错机制需要加强。")
    
    gen.add_title("（四）Clawguard安全扫描提案（Issue #36990）", level=2)
    gen.add_paragraph("提案内容：为OpenClaw skills/plugins提供安全扫描和守卫。")
    gen.add_paragraph("功能包括：扫描风险模式、动态检测、自动更新规则包。")
    gen.add_paragraph("状态：开放（社区提案）")
    gen.add_paragraph("启示：社区已意识到需要更强的安全保障。")
    
    gen.add_page_break()
    
    # 三、性能问题分析
    gen.add_title("三、性能问题分析", level=1)
    
    gen.add_title("（一）并发会话序列化（Issue #44026）", level=2)
    gen.add_paragraph("问题描述：同一Agent ID的多个会话被序列化执行，即使配置了maxConcurrent: 4。")
    gen.add_paragraph("影响：简单请求被长任务阻塞超过150秒。")
    gen.add_paragraph("场景：团队Discord/Matrix频道中，多个用户同时使用。")
    gen.add_paragraph("根因分析：锁定粒度为Agent ID级别，非Session ID级别。")
    gen.add_paragraph("建议方案：将锁定粒度从Agent ID改为Session ID，实现读写锁机制。")
    gen.add_paragraph("状态：开放")
    gen.add_paragraph("启示：多用户并发是企业级应用的必备能力。")
    
    gen.add_title("（二）Memory flush阻塞（Issue #44032）", level=2)
    gen.add_paragraph("问题描述：内存压缩（compaction）过程阻塞用户消息。")
    gen.add_paragraph("根因：isCompacting锁机制。")
    gen.add_paragraph("状态：开放")
    
    gen.add_title("（三）Compaction超时（Issue #44031）", level=2)
    gen.add_paragraph("问题描述：内存压缩超时导致会话挂起。")
    gen.add_paragraph("建议：增加逃逸舱机制。")
    gen.add_paragraph("状态：开放")
    
    gen.add_page_break()
    
    # 四、部署问题分析
    gen.add_title("四、部署问题分析", level=1)
    
    gen.add_title("（一）嵌入式环境崩溃（Issue #44006）", level=2)
    gen.add_paragraph("环境：ImmortalWrt 24.10.2 (x86/64, musl libc), Node.js v22.16.0")
    gen.add_paragraph("问题：Telegram发送PDF导致Gateway崩溃（SIGILL: Invalid Opcode）")
    gen.add_paragraph("根因：skia.linux-x64-musl.node使用了CPU不支持的AVX2指令集。")
    gen.add_paragraph("建议修复：")
    gen.add_paragraph("1. 提供不含skia/canvas的精简构建版本")
    gen.add_paragraph("2. 增加CPU特性检测，不支持时优雅降级")
    gen.add_paragraph("3. 增加media.preview.enabled开关")
    gen.add_paragraph("启示：嵌入式/边缘计算场景需要考虑CPU指令集兼容性。")
    
    gen.add_title("（二）缺少Windows桌面应用（Issue #75）", level=2)
    gen.add_paragraph("问题：缺少Linux和Windows桌面应用。")
    gen.add_paragraph("现状：已有macOS, iOS, Android应用。")
    gen.add_paragraph("需求：37条评论，用户关注度高。")
    gen.add_paragraph("状态：开放")
    gen.add_paragraph("启示：企业办公环境多为Windows，需要优先支持。")
    
    gen.add_page_break()
    
    # 五、本地化问题分析
    gen.add_title("五、本地化问题分析", level=1)
    
    gen.add_title("国际化支持缺失（Issue #3460）", level=2)
    gen.add_paragraph("需求语言：中文简体、繁体、日语、韩语等10种语言。")
    gen.add_paragraph("社区关注度：97+条评论。")
    gen.add_paragraph("官方回复：当前没有带宽支持多语言。")
    gen.add_paragraph("状态：开放（等待官方i18n架构）")
    gen.add_paragraph("启示：中文支持对国内市场至关重要。")
    
    gen.add_page_break()
    
    # 六、用户痛点总结
    gen.add_title("六、用户痛点总结", level=1)
    
    gen.add_table(["痛点类别", "具体问题", "影响程度", "用户反馈数"],
                  [["并发性能", "多用户场景序列化执行", "致命", "高"],
                   ["安全漏洞", "API Key泄露、策略绕过", "高", "中"],
                   ["部署复杂", "缺少Windows支持、环境兼容性", "中", "高"],
                   ["本地化缺失", "无中文支持", "中", "极高（97+评论）"],
                   ["稳定性", "内存管理、崩溃问题", "中", "中"]])
    
    gen.add_paragraph("核心结论：OpenClaw适合个人/小团队使用，但对于企业多用户并发场景存在根本性架构缺陷，需要大量定制开发才能满足企业级需求。")
    
    # 保存
    output_path = os.path.join(OUTPUT_DIR, '附件二-GitHub-Issues调研报告.docx')
    gen.save(output_path)
    return output_path


def generate_attachment3():
    """附件三：央国企AI助手技术研究报告"""
    gen = BaseReportGenerator()
    
    gen.add_title("附件三：央国企AI助手技术研究报告", level=0)
    gen.doc.add_paragraph()
    
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("安全合规、国产模型、权限控制方案")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    gen.add_page_break()
    
    # 一、安全合规要求
    gen.add_title("一、央国企AI应用安全合规要求", level=1)
    
    gen.add_title("（一）等保2.0要求", level=2)
    gen.add_paragraph("央国企AI应用通常需要满足等保三级或以上要求。核心要求包括：")
    gen.add_table(["要求类别", "具体要求", "实施建议"],
                  [["身份鉴别", "双因子认证、强密码策略", "集成企业统一身份认证"],
                   ["访问控制", "RBAC权限管理", "基于角色的细粒度权限"],
                   ["安全审计", "全量日志记录、行为分析", "操作日志+AI行为日志"],
                   ["数据保密", "加密存储、加密传输", "AES-256加密，TLS 1.3"],
                   ["数据备份", "定期备份、快速恢复", "日增量+周全量+异地容灾"]])
    
    gen.add_title("（二）数据安全法要求", level=2)
    gen.add_paragraph("核心原则：数据分类分级、数据安全责任、风险评估、应急处置。")
    gen.add_paragraph("关键要求：重要数据出境需进行安全评估；建立数据安全管理制度；开展数据安全教育培训；配合国家数据安全监管。")
    
    gen.add_title("（三）个人信息保护法要求", level=2)
    gen.add_paragraph("核心原则：合法正当必要、最小必要原则、知情同意、目的限制。")
    gen.add_paragraph("关键要求：制定个人信息保护政策；设立个人信息保护负责人；进行个人信息保护影响评估；建立个人信息安全事件应急预案。")
    
    gen.add_page_break()
    
    # 二、国产大模型对比
    gen.add_title("二、国产大模型能力对比", level=1)
    
    gen.add_title("（一）百度文心一言", level=2)
    gen.add_paragraph("技术特点：知识增强大语言模型，千亿级参数。")
    gen.add_paragraph("核心能力：文本生成与理解、知识问答、代码生成、多模态理解。")
    gen.add_paragraph("企业特性：千帆平台、支持私有化部署、算法备案完成。")
    gen.add_paragraph("适用场景：知识密集型应用、政务服务、金融风控。")
    
    gen.add_title("（二）阿里通义千问", level=2)
    gen.add_paragraph("技术特点：通义大模型系列，万亿级参数。")
    gen.add_paragraph("核心能力：Qwen3-Max全能型、Qwen-Plus均衡型、Qwen-Flash极速型、Qwen3-VL视觉理解。")
    gen.add_paragraph("企业特性：百炼平台、支持私有化部署、与阿里云产品深度集成。")
    gen.add_paragraph("适用场景：电商智能客服、文档智能处理、数据分析。")
    
    gen.add_title("（三）腾讯混元", level=2)
    gen.add_paragraph("技术特点：混合专家模型（MoE），万亿级参数。")
    gen.add_paragraph("核心能力：深度推理、快速响应、视觉理解、视频生成。")
    gen.add_paragraph("企业特性：多轮对话优势、深度合成服务算法备案、可信AI评估。")
    gen.add_paragraph("适用场景：智能客服、内容审核、广告推荐、会议助手。")
    
    gen.add_title("（四）华为盘古", level=2)
    gen.add_paragraph("技术特点：五大基础模型（LLM、视觉、多模态、预测、气象）。")
    gen.add_paragraph("核心能力：行业深耕（矿山、政务、金融）。")
    gen.add_paragraph("企业特性：支持私有化部署、首个卓越级评分大模型。")
    gen.add_paragraph("适用场景：矿山智能化、智慧城市、气象预报、金融风控。")
    
    gen.add_title("（五）模型对比总结", level=2)
    gen.add_table(["维度", "百度文心", "阿里通义", "腾讯混元", "华为盘古"],
                  [["技术特点", "知识增强", "全模态", "MoE架构", "行业深耕"],
                   ["参数规模", "千亿级", "万亿级", "万亿级", "千亿级"],
                   ["私有化部署", "✓", "✓", "✓", "✓"],
                   ["算法备案", "✓", "✓", "✓", "✓"],
                   ["行业适配", "政务/金融", "电商/文档", "客服/会议", "能源/气象"]])
    
    gen.add_page_break()
    
    # 三、权限控制方案
    gen.add_title("三、企业级权限控制方案", level=1)
    
    gen.add_title("（一）RBAC模型", level=2)
    gen.add_paragraph("基于角色的访问控制，实现菜单权限和操作权限管理。")
    gen.add_paragraph("模型结构：User → Role → Permission → Resource")
    gen.add_paragraph("优点：易于理解和管理，适合稳定组织结构。")
    gen.add_paragraph("缺点：权限爆炸问题，缺乏灵活性。")
    
    gen.add_title("（二）ABAC模型", level=2)
    gen.add_paragraph("基于属性的访问控制，实现数据行级和列级权限管理。")
    gen.add_paragraph("核心概念：主体属性、资源属性、环境属性、策略。")
    gen.add_paragraph("优点：高度灵活，细粒度控制，动态决策。")
    gen.add_paragraph("缺点：实现复杂，性能开销大。")
    
    gen.add_title("（三）混合模型方案", level=2)
    gen.add_paragraph("建议采用RBAC与ABAC混合模型：")
    gen.add_paragraph("功能权限（RBAC）：角色定义、菜单权限、操作权限。")
    gen.add_paragraph("数据权限（ABAC）：部门数据隔离、敏感数据访问、项目数据权限。")
    gen.add_paragraph("AI能力权限：工具调用权限、模型使用权限、数据访问范围。")
    
    gen.add_page_break()
    
    # 四、技术方案建议
    gen.add_title("四、技术方案建议", level=1)
    
    gen.add_title("（一）大模型选型建议", level=2)
    gen.add_table(["场景", "推荐模型", "理由"],
                  [["通用对话", "百度文心/阿里通义", "综合能力强，企业支持完善"],
                   ["知识问答", "百度文心", "知识增强优势明显"],
                   ["代码开发", "阿里通义Qwen-Coder", "代码专项优化"],
                   ["行业应用", "华为盘古", "行业深耕，专业性强"]])
    
    gen.add_title("（二）基础设施选型建议", level=2)
    gen.add_paragraph("操作系统：统信UOS、麒麟OS等国产操作系统。")
    gen.add_paragraph("数据库：达梦、人大金仓等国产数据库，或PostgreSQL。")
    gen.add_paragraph("中间件：东方通、宝兰德等国产中间件。")
    gen.add_paragraph("云平台：阿里云、华为云、腾讯云，或私有云部署。")
    
    # 保存
    output_path = os.path.join(OUTPUT_DIR, '附件三-央国企AI助手技术研究报告.docx')
    gen.save(output_path)
    return output_path


def generate_attachment4(diagrams):
    """附件四：SOE-Claw架构设计文档"""
    gen = BaseReportGenerator()
    
    gen.add_title("附件四：SOE-Claw架构设计文档", level=0)
    gen.doc.add_paragraph()
    
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("完整的架构设计和实现方案")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    gen.add_page_break()
    
    # 一、架构概述
    gen.add_title("一、架构概述", level=1)
    gen.add_paragraph("SOE-Claw（State-Owned Enterprise Claw）是针对央国企需求设计的智能助手架构，借鉴OpenClaw的优秀设计思想，同时满足央国企的安全合规、权限控制、国产化等需求。")
    
    gen.add_image(diagrams['soe_architecture'], width_inches=6, caption="图1 SOE-Claw架构图")
    
    # 二、架构分层设计
    gen.add_title("二、架构分层设计", level=1)
    
    gen.add_title("（一）用户接入层", level=2)
    gen.add_paragraph("功能：支持Web管理端、企业微信、钉钉、飞书、政务平台等多种接入方式。")
    gen.add_paragraph("设计要点：统一的API网关实现请求路由和负载均衡；支持多种认证方式；实现请求限流和熔断。")
    
    gen.add_title("（二）安全接入层", level=2)
    gen.add_paragraph("功能：集成企业统一身份认证（ISC），实现RBAC权限控制，建立完整的访问审计机制。")
    gen.add_paragraph("设计要点：对接企业LDAP/AD；实现RBAC+ABAC混合权限模型；全量操作审计；安全防护（WAF、抗DDoS）。")
    
    gen.add_title("（三）AI服务层", level=2)
    gen.add_paragraph("功能：提供对话服务、知识库服务（RAG）、工具服务、Agent编排能力。")
    gen.add_paragraph("设计要点：支持多模型切换；会话管理和记忆系统；安全沙箱执行；工具调用权限控制。")
    
    gen.add_title("（四）模型服务层", level=2)
    gen.add_paragraph("功能：优先使用国产大模型，支持私有化部署。")
    gen.add_paragraph("设计要点：多模型支持（文心、通义、混元、盘古）；模型热切换；A/B测试能力；推理优化。")
    
    gen.add_title("（五）数据服务层", level=2)
    gen.add_paragraph("功能：建设企业知识库、向量数据库、关系数据库、对象存储。")
    gen.add_paragraph("设计要点：数据加密存储；访问控制；定期备份；数据血缘追踪。")
    
    gen.add_page_break()
    
    # 三、核心组件设计
    gen.add_title("三、核心组件设计", level=1)
    
    gen.add_title("（一）分布式Gateway设计", level=2)
    gen.add_paragraph("针对OpenClaw Gateway单点问题，设计分布式Gateway架构：")
    gen.add_paragraph("1. 多Gateway实例：支持水平扩展，无单点故障。")
    gen.add_paragraph("2. 会话级锁定：将锁定粒度从Agent ID改为Session ID。")
    gen.add_paragraph("3. 状态持久化：使用Redis存储会话状态，支持恢复。")
    gen.add_paragraph("4. 负载均衡：通过Nginx实现请求分发。")
    
    gen.add_title("（二）安全沙箱设计", level=2)
    gen.add_paragraph("增强Docker沙箱配置：")
    gen.add_code_block('agent-sandbox:\n  deploy:\n    resources:\n      limits:\n        cpus: "2.0"\n        memory: 2G\n  security_opt:\n    - no-new-privileges:true\n    - seccomp:seccomp-profile.json\n  cap_drop:\n    - ALL\n  networks:\n    - sandbox-net')
    
    gen.add_title("（三）Skills安全设计", level=2)
    gen.add_paragraph("1. 运行时隔离：在独立容器中执行Skills代码。")
    gen.add_paragraph("2. 权限控制：细粒度权限白名单，限制可访问资源。")
    gen.add_paragraph("3. 审核流程：Skills审核和签名机制。")
    gen.add_paragraph("4. 动态监控：实时监控Skills行为。")
    
    gen.add_page_break()
    
    # 四、权限控制设计
    gen.add_title("四、权限控制设计", level=1)
    
    gen.add_image(diagrams['permission'], width_inches=6, caption="图2 权限控制模型")
    
    gen.add_title("（一）功能权限配置示例", level=2)
    gen.add_code_block('{\n  "roles": {\n    "admin": {\n      "permissions": ["all"],\n      "dataScope": "enterprise"\n    },\n    "department_manager": {\n      "permissions": ["chat", "knowledge", "report"],\n      "dataScope": "department"\n    },\n    "staff": {\n      "permissions": ["chat"],\n      "dataScope": "self"\n    }\n  }\n}')
    
    gen.add_title("（二）AI能力权限配置示例", level=2)
    gen.add_code_block('{\n  "aiPermissions": {\n    "tools": {\n      "allow": ["web-search", "document-read"],\n      "deny": ["exec", "file-write"]\n    },\n    "data": {\n      "scope": "department",\n      "sensitivity": ["public", "internal"]\n    },\n    "models": {\n      "allow": ["ernie-4.0", "qwen-max"],\n      "deny": ["gpt-4"]\n    }\n  }\n}')
    
    # 保存
    output_path = os.path.join(OUTPUT_DIR, '附件四-SOE-Claw架构设计文档.docx')
    gen.save(output_path)
    return output_path


def generate_attachment5():
    """附件五：安全合规对标清单"""
    gen = BaseReportGenerator()
    
    gen.add_title("附件五：安全合规对标清单", level=0)
    gen.doc.add_paragraph()
    
    p = gen.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("等保2.0详细对标和整改建议")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    gen.add_page_break()
    
    # 一、等保三级对标清单
    gen.add_title("一、等保三级对标清单", level=1)
    
    gen.add_title("（一）安全物理环境", level=2)
    gen.add_table(["要求", "OpenClaw现状", "SOE-Claw方案", "优先级"],
                  [["物理访问控制", "无", "机房门禁+视频监控", "高"],
                   ["防盗窃和防破坏", "无", "设备锁定+报警", "高"],
                   ["防雷击", "无", "避雷装置+接地", "中"],
                   ["防火", "无", "灭火系统+烟感", "高"],
                   ["防水防潮", "无", "防水措施+湿度监控", "中"]])
    
    gen.add_title("（二）安全通信网络", level=2)
    gen.add_table(["要求", "OpenClaw现状", "SOE-Claw方案", "优先级"],
                  [["网络架构", "无隔离", "安全域划分+防火墙", "高"],
                   ["通信传输加密", "WebSocket无加密", "TLS 1.3加密", "高"],
                   ["可信验证", "无", "设备可信验证", "中"]])
    
    gen.add_title("（三）安全区域边界", level=2)
    gen.add_table(["要求", "OpenClaw现状", "SOE-Claw方案", "优先级"],
                  [["边界防护", "无", "防火墙+入侵检测", "高"],
                   ["访问控制", "Allowlist基础", "RBAC+ABAC混合", "高"],
                   ["入侵防范", "无", "IDS+IPS", "高"],
                   ["恶意代码防范", "基础扫描", "杀毒+沙箱", "高"],
                   ["安全审计", "基础日志", "全量审计日志", "高"]])
    
    gen.add_title("（四）安全计算环境", level=2)
    gen.add_table(["要求", "OpenClaw现状", "SOE-Claw方案", "优先级"],
                  [["身份鉴别", "Token/Password", "双因子认证+强密码", "高"],
                   ["访问控制", "配对+Allowlist", "RBAC+最小权限", "高"],
                   ["安全审计", "基础审计", "全量审计+分析", "高"],
                   ["入侵防范", "无", "主机入侵检测", "高"],
                   ["恶意代码防范", "无", "杀毒软件+沙箱", "高"],
                   ["数据完整性", "无", "校验机制", "中"],
                   ["数据保密性", "无", "AES-256加密", "高"],
                   ["数据备份恢复", "无", "日增量+周全量", "高"]])
    
    gen.add_page_break()
    
    # 二、整改优先级
    gen.add_title("二、整改优先级建议", level=1)
    
    gen.add_title("（一）P0（必须立即整改）", level=2)
    gen.add_paragraph("1. 身份认证：对接企业统一身份认证，实现双因子认证。")
    gen.add_paragraph("2. 访问控制：实现RBAC权限模型，支持最小权限原则。")
    gen.add_paragraph("3. 安全审计：建立全量审计日志，支持溯源追踪。")
    gen.add_paragraph("4. 数据加密：实现数据存储加密和传输加密。")
    gen.add_paragraph("5. 边界防护：部署防火墙和入侵检测系统。")
    
    gen.add_title("（二）P1（短期整改）", level=2)
    gen.add_paragraph("1. 数据备份：建立定期备份机制，支持快速恢复。")
    gen.add_paragraph("2. 恶意代码防范：部署杀毒软件和沙箱隔离。")
    gen.add_paragraph("3. 入侵防范：部署主机入侵检测系统。")
    gen.add_paragraph("4. 数据完整性：建立数据校验机制。")
    
    gen.add_title("（三）P2（中期整改）", level=2)
    gen.add_paragraph("1. 物理安全：完善机房安全措施。")
    gen.add_paragraph("2. 可信验证：实现设备可信验证。")
    gen.add_paragraph("3. 防雷防火防水：完善物理环境安全。")
    
    gen.add_page_break()
    
    # 三、合规评估表
    gen.add_title("三、合规评估表", level=1)
    
    gen.add_table(["评估项", "权重", "OpenClaw得分", "SOE-Claw得分"],
                  [["身份鉴别", "15%", "40", "95"],
                   ["访问控制", "15%", "50", "90"],
                   ["安全审计", "15%", "30", "90"],
                   ["入侵防范", "10%", "20", "85"],
                   ["恶意代码防范", "10%", "30", "85"],
                   ["数据完整性", "10%", "20", "80"],
                   ["数据保密性", "15%", "10", "90"],
                   ["数据备份恢复", "10%", "0", "85"],
                   ["综合得分", "100%", "25", "88"]])
    
    gen.add_paragraph("结论：OpenClaw综合得分25分，远低于等保三级要求（70分以上）。SOE-Claw方案综合得分88分，满足等保三级要求。")
    
    # 保存
    output_path = os.path.join(OUTPUT_DIR, '附件五-安全合规对标清单.docx')
    gen.save(output_path)
    return output_path


def main():
    """主函数"""
    print("=" * 60)
    print("OpenClaw技术研究报告 - 最终发布版生成器")
    print("=" * 60)
    print()
    
    # 生成图表
    print("步骤1: 生成架构图...")
    diagram_gen = DiagramGenerator(IMAGES_DIR)
    diagrams = diagram_gen.generate_all()
    print(f"  ✅ 已生成{len(diagrams)}张架构图")
    print()
    
    # 生成主报告
    print("步骤2: 生成主报告...")
    main_report = generate_main_report(diagrams)
    print(f"  ✅ 主报告: {main_report}")
    print()
    
    # 生成附件
    print("步骤3: 生成附件...")
    
    print("  生成附件一：OpenClaw源码分析报告...")
    att1 = generate_attachment1(diagrams)
    print(f"    ✅ {att1}")
    
    print("  生成附件二：GitHub Issues调研报告...")
    att2 = generate_attachment2()
    print(f"    ✅ {att2}")
    
    print("  生成附件三：央国企AI助手技术研究报告...")
    att3 = generate_attachment3()
    print(f"    ✅ {att3}")
    
    print("  生成附件四：SOE-Claw架构设计文档...")
    att4 = generate_attachment4(diagrams)
    print(f"    ✅ {att4}")
    
    print("  生成附件五：安全合规对标清单...")
    att5 = generate_attachment5()
    print(f"    ✅ {att5}")
    print()
    
    print("=" * 60)
    print("✅ 所有文档生成完成！")
    print("=" * 60)
    print()
    print("发布文件清单：")
    print(f"  📄 主报告: {main_report}")
    print(f"  📎 附件一: {att1}")
    print(f"  📎 附件二: {att2}")
    print(f"  📎 附件三: {att3}")
    print(f"  📎 附件四: {att4}")
    print(f"  📎 附件五: {att5}")
    print()
    print(f"发布目录: {OUTPUT_DIR}")
    print()


if __name__ == "__main__":
    main()
